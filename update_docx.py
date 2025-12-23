# update_docx.py
from docx import Document
import io

def update_questionnaire_with_answers(original_doc_bytes, answers_json):
    """
    Updates answers in the questionnaire.
    If table exists, updates the table.
    If no table, creates a new table with questions and answers.
    """

    doc = Document(io.BytesIO(original_doc_bytes))

    answers_map = {
        a["question_index"]: a
        for a in answers_json.get("answers", [])
    }

    if doc.tables:
        # Update existing table
        table = doc.tables[0]

        # Ensure Answer column exists
        current_cols = len(table.columns)
        if current_cols < 4:
            table.add_column(table.columns[0].width)
            table.rows[0].cells[3].text = "Answer"

        question_index = 0
        for row_idx, row in enumerate(table.rows):
            if row_idx == 0:
                continue  # skip header row

            question_index += 1
            answer_data = answers_map.get(question_index)
            if answer_data:
                answer = str(answer_data["answer"])
            else:
                answer = "Not found in source"

            if len(row.cells) > 3:
                row.cells[3].text = answer
    else:
        # No table, create a new table with all questions and answers
        table = doc.add_table(rows=1, cols=4)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Question'
        hdr_cells[1].text = 'Description'
        hdr_cells[2].text = 'Response Type'
        hdr_cells[3].text = 'Answer'

        for ans in answers_json.get("answers", []):
            row_cells = table.add_row().cells
            row_cells[0].text = str(ans.get("question", ""))
            row_cells[1].text = str(ans.get("description", ""))
            row_cells[2].text = str(ans.get("response_type", ""))
            row_cells[3].text = str(ans.get("answer", ""))

    out_stream = io.BytesIO()
    doc.save(out_stream)
    return out_stream.getvalue()
