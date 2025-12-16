# list_sharepoint_files.py
import os
from dotenv import load_dotenv
from onedrive_client import OneDriveClient
from docx import Document

load_dotenv()

BASE_URL = os.getenv("BASE_URL")
FEDAUTH = os.getenv("FEDAUTH")
RTFA = os.getenv("RTFA")

QUESTION_PATH = os.getenv("QUESTION_PATH")
EVIDENCE_PATH = os.getenv("EVIDENCE_PATH")

def count_questions_in_docx_table(docx_path):
    doc = Document(docx_path)

    if not doc.tables:
        return 0

    table = doc.tables[0]
    count = 0

    for i, row in enumerate(table.rows):
        if i == 0:
            continue  # header

        if row.cells[0].text.strip():
            count += 1

    return count

def list_files():
    client = OneDriveClient(BASE_URL, FEDAUTH, RTFA)

    print("\n📄 QUESTIONNAIRE FILE\n")

    local_q_file = "questionnaire_temp.docx"
    client.download_file(QUESTION_PATH, local_q_file)

    question_name = QUESTION_PATH.split("/")[-1]
    question_count = count_questions_in_docx_table(local_q_file)

    print(f"• {question_name}")
    print(f"  ServerRelativeUrl: {QUESTION_PATH}")
    print(f"  ✅ Number of questions (from table): {question_count}")

    print("\n📂 EVIDENCE FILES\n")

    files = client.list_files(EVIDENCE_PATH)

    if not files:
        print("No evidence files found.")
        return

    for i, file_url in enumerate(files, start=1):
        filename = file_url.split("/")[-1]
        print(f"{i}. {filename}")
        print(f"   ServerRelativeUrl: {file_url}")

if __name__ == "__main__":
    list_files()
