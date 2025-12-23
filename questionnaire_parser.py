#questionnaire_parser.py
from docx import Document
import io
import re

def extract_questions_from_table(docx_bytes):
    """
    Extracts structured questions from a table-based questionnaire DOCX.

    Expected table format:
    | Question | Description / Clarification | Response Type |
    """

    doc = Document(io.BytesIO(docx_bytes))

    if not doc.tables:
        raise Exception("No tables found in questionnaire")

    table = doc.tables[0]  # assume first table contains questions

    questions = []

    for row_index, row in enumerate(table.rows):
        if row_index == 0:
            continue  # skip header row

        question = row.cells[0].text.strip()
        description = row.cells[1].text.strip() if len(row.cells) > 1 else ""
        response_type = row.cells[2].text.strip() if len(row.cells) > 2 else ""

        if question:
            questions.append({
                "question_index": len(questions) + 1,
                "question": question,
                "description": description,
                "response_type": response_type
            })

    return questions

def extract_questions_from_text(docx_bytes):
    """
    Extracts questions from plain text in the DOCX if no table is found.
    Assumes questions are lines ending with '?' or containing question words.
    """
    doc = Document(io.BytesIO(docx_bytes))
    
    full_text = ""
    for para in doc.paragraphs:
        full_text += para.text + "\n"
    
    # Split into potential questions by lines
    lines = [line.strip() for line in full_text.split('\n') if line.strip()]
    
    questions = []
    for line in lines:
        # Consider it a question if it ends with ? or contains question words
        if '?' in line or any(word in line.lower() for word in ['what', 'how', 'why', 'when', 'where', 'who', 'which', 'do you', 'does', 'is', 'are', 'can', 'could', 'should', 'would']):
            questions.append({
                "question_index": len(questions) + 1,
                "question": line,
                "description": "",
                "response_type": ""
            })
    
    return questions

def extract_questions(docx_bytes):
    """
    Tries to extract questions from table first, falls back to text parsing.
    """
    try:
        return extract_questions_from_table(docx_bytes)
    except Exception:
        return extract_questions_from_text(docx_bytes)
